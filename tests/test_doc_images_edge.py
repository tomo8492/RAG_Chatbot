"""文書内画像まわりの徹底(エッジケース)テスト。

ノイズ除外(ロゴ・アイコン)、取り込み上限、形式変換(BMP/パレットPNG/巨大画像)、
壊れ入力への耐性、Word表中の画像、PDFの繰り返しロゴ、増分再構築でのリンク維持、
WebP参考図のエクスポート互換まで、実装した画像機能の境界を固定する。
Ollama・埋め込みはスタブ(ChromaDB・ファイル生成は実物)。pytest/単体実行両対応。
"""
import contextlib
import io
import os
import random
import shutil
import sys
import tempfile
import zipfile
from pathlib import Path

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from app import db, doc_images, export, ocr, rag   # noqa: E402
from app.config import settings                     # noqa: E402


def _has(mod: str) -> bool:
    try:
        __import__(mod)
        return True
    except Exception:
        return False


@contextlib.contextmanager
def patched(obj, name, value):
    old = getattr(obj, name)
    setattr(obj, name, value)
    try:
        yield
    finally:
        setattr(obj, name, old)


@contextlib.contextmanager
def temp_env():
    d = Path(tempfile.mkdtemp(prefix="docimg_edge_")).resolve()
    (d / "data").mkdir()
    docs = d / "docs"
    docs.mkdir()
    old = (settings.data_dir, settings.chroma_dir, settings.db_path)
    settings.data_dir = d / "data"
    settings.chroma_dir = d / "data" / "chroma"
    settings.db_path = d / "data" / "t.db"
    old_client = rag._client
    rag._client = None
    try:
        db.init_db()
        yield docs
    finally:
        rag._client = old_client
        settings.data_dir, settings.chroma_dir, settings.db_path = old
        shutil.rmtree(d, ignore_errors=True)


class FakeEmb:
    def embed_query(self, t):
        return [0.1] * 8

    def embed_documents(self, ts):
        return [[0.1] * 8 for _ in ts]


def _noise_img(w: int, h: int, seed: int = 0):
    from PIL import Image
    raw = random.Random(seed).randbytes(w * h * 3)
    return Image.frombytes("RGB", (w, h), raw)


def _png_bytes(w: int = 240, h: int = 160, seed: int = 0) -> bytes:
    buf = io.BytesIO()
    _noise_img(w, h, seed).save(buf, format="PNG")
    return buf.getvalue()


def _xlsx_with_images(path: Path, images: list[tuple[str, Path]]):
    """images = [(シート名, 画像パス), ...]。同名シートには複数貼る。"""
    import openpyxl
    from openpyxl.drawing.image import Image as XLImage
    wb = openpyxl.Workbook()
    sheets: dict = {}
    anchor_row: dict = {}
    for sheet, img in images:
        if sheet not in sheets:
            ws = wb.active if not sheets else wb.create_sheet(sheet)
            ws.title = sheet
            ws["A1"] = f"{sheet} の説明テキスト"
            sheets[sheet] = ws
            anchor_row[sheet] = 2
        ws = sheets[sheet]
        ws.add_image(XLImage(str(img)), f"C{anchor_row[sheet]}")
        anchor_row[sheet] += 12
    wb.save(str(path))


# ---------------- ノイズ除外(ロゴ・繰り返し) ----------------
def test_same_image_on_many_sheets_treated_as_logo():
    with temp_env() as docs:
        png = docs / "logo.png"
        png.write_bytes(_png_bytes(seed=1))
        _xlsx_with_images(docs / "a.xlsx",
                          [("S1", png), ("S2", png), ("S3", png), ("S4", png), ("S5", png)])
        loc_map, images = doc_images.extract_for_file("idxL", docs / "a.xlsx")
        # 同一画像は MAX_REPEAT_LOCS(3) を超える場所には紐付けない(ロゴ判定)
        assert len(loc_map) <= doc_images.MAX_REPEAT_LOCS
        assert len(images) == 1                                  # 物理・説明対象は1枚
        files = list((settings.data_dir / "doc_images" / "idxL").glob("*"))
        assert len(files) == 1


def test_per_loc_cap_limits_images_per_sheet():
    with temp_env() as docs:
        paths = []
        for i in range(10):
            p = docs / f"img{i}.png"
            p.write_bytes(_png_bytes(100, 80, seed=10 + i))
            paths.append(("S1", p))
        _xlsx_with_images(docs / "many.xlsx", paths)
        loc_map, images = doc_images.extract_for_file("idxC", docs / "many.xlsx")
        assert len(loc_map.get("シート:S1", [])) <= doc_images.MAX_PER_LOC
        assert len(images) == 10                                  # 保存・説明対象自体は10枚


def test_per_file_cap():
    with temp_env() as docs:
        entries = []
        for i in range(45):
            p = docs / f"i{i}.png"
            p.write_bytes(_png_bytes(100, 80, seed=100 + i))
            entries.append((f"S{i % 9}", p))
        _xlsx_with_images(docs / "huge.xlsx", entries)
        _, images = doc_images.extract_for_file("idxF", docs / "huge.xlsx")
        assert len(images) <= doc_images.MAX_PER_FILE


# ---------------- 形式変換・寸法 ----------------
def test_bmp_converted_to_web_format():
    buf = io.BytesIO()
    _noise_img(200, 150, seed=2).save(buf, format="BMP")
    out = doc_images._normalize(buf.getvalue())
    assert out is not None and out[1] in ("png", "jpg")


def test_oversize_image_shrunk_to_max_dim():
    out = doc_images._normalize(_png_bytes(2200, 900, seed=3))
    assert out is not None
    from PIL import Image
    img = Image.open(io.BytesIO(out[0]))
    assert max(img.size) <= doc_images.MAX_DIM


def test_palette_png_oversize_keeps_alpha_path():
    from PIL import Image
    big = _noise_img(1800, 600, seed=4).convert("P", palette=Image.ADAPTIVE)
    buf = io.BytesIO()
    big.save(buf, format="PNG")
    out = doc_images._normalize(buf.getvalue())
    assert out is not None and out[1] == "png"
    assert Image.open(io.BytesIO(out[0])).size[0] <= doc_images.MAX_DIM


def test_garbage_and_fake_emf_skipped():
    assert doc_images._normalize(random.Random(5).randbytes(8000)) is None   # 形式不明
    fake_emf = b"\x01\x00\x00\x00" + b"\x00" * 36 + b" EMF" + b"\x00" * 8000
    assert doc_images._normalize(fake_emf) is None                # PILで開けないEMFは除外
    assert doc_images.extract_for_file("x", Path("/no/such.txt")) == ({}, [])   # 未対応拡張子


# ---------------- Word: 表のセル内の画像 ----------------
def test_docx_image_inside_table_linked_to_heading():
    with temp_env() as docs:
        from docx import Document
        png = docs / "fig.png"
        png.write_bytes(_png_bytes(seed=6))
        d = Document()
        d.add_heading("設置手順", level=1)
        tbl = d.add_table(rows=1, cols=2)
        tbl.rows[0].cells[0].text = "手順1: 図のとおり配線する"
        tbl.rows[0].cells[1].paragraphs[0].add_run().add_picture(str(png))
        d.save(str(docs / "m.docx"))
        loc_map, images = doc_images.extract_for_file("idxW", docs / "m.docx")
        assert loc_map.get("設置手順"), "表中の画像が直前の見出しに紐付くはず"
        assert len(images) == 1


# ---------------- PDF: 全ページ繰り返しのロゴ除外 ----------------
def test_pdf_repeated_logo_excluded_unique_kept():
    if not _has("fitz"):
        return   # pymupdf 無し環境は自動スキップ
    with temp_env() as docs:
        import fitz
        logo = docs / "logo.jpg"
        _noise_img(200, 100, seed=7).save(str(logo), format="JPEG")
        uniq = docs / "uniq.jpg"
        _noise_img(400, 300, seed=8).save(str(uniq), format="JPEG")
        pdf = fitz.open()
        xref = 0
        for i in range(5):
            page = pdf.new_page()
            page.insert_text((72, 60), f"ページ{i + 1}")
            if i == 0:
                xref = page.insert_image(fitz.Rect(72, 80, 272, 180), filename=str(logo))
                page.insert_image(fitz.Rect(72, 200, 472, 500), filename=str(uniq))
            else:
                page.insert_image(fitz.Rect(72, 80, 272, 180), xref=xref)   # 同一xref=ロゴ
        pdf.save(str(docs / "doc.pdf"))
        pdf.close()
        loc_map, images = doc_images.extract_for_file("idxP", docs / "doc.pdf")
        assert len(images) == 1                      # 残るのはユニーク画像だけ
        assert list(loc_map) == ["p.1"]


# ---------------- 増分再構築: リンク維持・重複なし ----------------
def test_rebuild_unchanged_keeps_links_and_no_duplicate_files():
    with temp_env() as docs:
        png = docs / "fig.png"
        png.write_bytes(_png_bytes(seed=9))
        _xlsx_with_images(docs / "手順.xlsx", [("組立", png)])
        idx = db.create_index("資料", [str(docs)])
        with patched(rag, "get_embedder", lambda: FakeEmb()), \
             patched(ocr, "describe_image_png", lambda b: "組立図"):
            r1 = rag.build_index(idx["id"], [str(docs)])
            col = rag._collection(f"idx_{idx['id']}")
            n1 = col.count()
            files1 = sorted((settings.data_dir / "doc_images" / idx["id"]).glob("*"))
            r2 = rag.build_index(idx["id"], [str(docs)])     # 変更なしで再構築
            n2 = col.count()
            files2 = sorted((settings.data_dir / "doc_images" / idx["id"]).glob("*"))
        assert r1["status"] == r2["status"] == "ready"
        assert n1 == n2, "据置でチャンクが増減しない"
        assert files1 == files2 and len(files1) == 1, "画像ファイルが重複保存されない"
        got = col.get(include=["metadatas", "documents"])
        assert any(m.get("images") for m in got["metadatas"]), "再構築後もリンク維持"
        assert any(d.startswith("〔図〕") for d in got["documents"])


# ---------------- エクスポート: WebP参考図のPNG変換 ----------------
def test_export_webp_figure_converted_and_embedded():
    import base64
    buf = io.BytesIO()
    _noise_img(240, 160, seed=11).save(buf, format="WEBP")
    figs = [{"data": base64.b64encode(buf.getvalue()).decode(), "caption": "WebPの図"}]
    html = export.to_html("# t\n本文", figures=figs).decode("utf-8")
    assert "data:image/png;base64," in html           # PNGへ変換されて埋め込まれる
    data = export.to_docx("# t\n本文", figures=figs)
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        assert any(n.startswith("word/media/") for n in z.namelist())
    data = export.to_pptx("# t\n本文", figures=figs)
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        assert any(n.startswith("ppt/media/") for n in z.namelist())


if __name__ == "__main__":
    fns = [v for k, v in sorted(globals().items()) if k.startswith("test_") and callable(v)]
    passed = 0
    for fn in fns:
        try:
            fn()
            passed += 1
        except AssertionError as e:
            print(f"FAIL {fn.__name__}: {e}")
        except Exception as e:  # noqa: BLE001
            print(f"ERROR {fn.__name__}: {e!r}")
    print(f"{passed}/{len(fns)} passed")
    sys.exit(0 if passed == len(fns) else 1)

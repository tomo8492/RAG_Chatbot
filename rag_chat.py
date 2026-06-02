#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
rag_chat.py
フォルダを選択し、中のファイル(PDF/Word/Excel)をもとに回答する
ローカルRAGチャットボット(GUI / Claude風)。

機能:
  - フォルダ選択ボタン / チャットウィンドウ / 質問入力欄
  - 回答生成中の「停止」ボタン(生成途中で中断できる)
  - すべての処理ログとエラーを「エディタのターミナル」に出力
  - Ollama + qwen3-32b:latest で完全ローカル動作(社外送信ゼロ)
  - 回答に出典ファイル名を明示

実行(ターミナルから):
  py -3.12 rag_chat.py
"""

import sys
import time
import threading
import queue
import traceback
from pathlib import Path

import tkinter as tk
from tkinter import filedialog, font as tkfont, messagebox


# ============================================================
#  ログ出力(必ずターミナルに出る)
# ============================================================
def log(msg, level="INFO"):
    ts = time.strftime("%H:%M:%S")
    print(f"[{ts}] [{level}] {msg}", flush=True)


def log_exc(where):
    print(f"\n{'=' * 60}", flush=True)
    print(f"[ERROR] {where} で例外が発生しました:", flush=True)
    traceback.print_exc()
    print(f"{'=' * 60}\n", flush=True)


# ============================================================
#  カラーパレット (Claude風)
# ============================================================
COL = {
    "bg": "#F0EEE6", "panel": "#FAF9F5", "card": "#FFFFFF",
    "ink": "#2B2A27", "ink_soft": "#6E6B63",
    "accent": "#D97757", "accent_dk": "#C25B3C",
    "stop": "#B5483A", "stop_dk": "#963729",
    "line": "#E3DFD3", "green": "#3D7A4E", "btn_sub": "#EDEAE0",
}


# ============================================================
#  文書ローダ(形式別)
# ============================================================
def load_pdf(path: Path):
    import fitz  # PyMuPDF
    docs = []
    with fitz.open(path) as pdf:
        for i, page in enumerate(pdf):
            text = page.get_text().strip()
            if text:
                docs.append({"text": text, "source": path.name, "loc": f"p.{i + 1}"})
    return docs


def load_docx(path: Path):
    from docx import Document
    doc = Document(path)
    docs, buf, head = [], [], "本文"
    for para in doc.paragraphs:
        t = para.text.strip()
        if not t:
            continue
        if para.style and para.style.name.startswith(("Heading", "見出し")):
            if buf:
                docs.append({"text": "\n".join(buf), "source": path.name, "loc": head})
                buf = []
            head = t
        buf.append(t)
    if buf:
        docs.append({"text": "\n".join(buf), "source": path.name, "loc": head})
    return docs


def load_xlsx(path: Path):
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    docs = []
    for ws in wb.worksheets:
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            docs.append({"text": "\n".join(rows), "source": path.name, "loc": f"シート:{ws.title}"})
    wb.close()
    return docs


LOADERS = {".pdf": load_pdf, ".docx": load_docx, ".xlsx": load_xlsx}

SYSTEM_PROMPT = """あなたは社内文書アシスタントです。以下のルールを厳守してください。

1. 回答は必ず【参考資料】の内容だけに基づいて答える。資料にない情報は推測せず「資料内に該当記載が見つかりません」と答える。
2. 回答の末尾に、必ず根拠とした出典を「出典: ファイル名 (場所)」の形式で列挙する。
3. 専門用語や条項はできるだけ原文のまま正確に引用する。
4. 不確かな場合は確実なことだけを述べ、断定しない。
"""

ANSWER_TEMPLATE = """{system}

【参考資料】
{context}

【質問】
{question}

【回答】"""


# ============================================================
#  GUIアプリ
# ============================================================
class RagApp:
    def __init__(self, root):
        self.root = root
        self.root.title("社内文書アシスタント")
        self.root.geometry("900x720")
        self.root.configure(bg=COL["bg"])
        self.root.minsize(720, 560)

        self.folder = None
        self.retriever = None
        self.llm = None
        self.embed_model = "intfloat/multilingual-e5-small"
        self.model_name = "qwen3-32b:latest"
        self.msg_queue = queue.Queue()

        # 停止制御用フラグ(回答生成中にTrueで中断要求)
        self.stop_flag = threading.Event()
        self.generating = False

        self._setup_fonts()
        self._build_ui()
        self.root.after(80, self._poll_queue)
        log("GUI起動完了。フォルダを選択してください。")

    def _setup_fonts(self):
        fam = "Yu Gothic UI"
        self.f_title = tkfont.Font(family=fam, size=15, weight="bold")
        self.f_body = tkfont.Font(family=fam, size=11)
        self.f_bold = tkfont.Font(family=fam, size=11, weight="bold")
        self.f_small = tkfont.Font(family=fam, size=9)

    def _btn(self, parent, text, command, primary=False):
        bg = COL["accent"] if primary else COL["btn_sub"]
        fg = "#FFFFFF" if primary else COL["ink"]
        active = COL["accent_dk"] if primary else COL["line"]
        return tk.Button(parent, text=text, command=command,
                         bg=bg, fg=fg, activebackground=active, activeforeground=fg,
                         font=self.f_bold, relief="flat", bd=0,
                         padx=16, pady=8, cursor="hand2")

    def _build_ui(self):
        header = tk.Frame(self.root, bg=COL["panel"], height=64)
        header.pack(fill="x")
        header.pack_propagate(False)
        tk.Frame(self.root, bg=COL["accent"], height=3).pack(fill="x")

        dot = tk.Canvas(header, width=30, height=30, bg=COL["panel"], highlightthickness=0)
        dot.create_oval(4, 4, 26, 26, fill=COL["accent"], outline="")
        dot.create_text(15, 15, text="\u2733", fill="#FFFFFF", font=("Yu Gothic UI", 12, "bold"))
        dot.pack(side="left", padx=(18, 8), pady=17)
        tk.Label(header, text="社内文書アシスタント", bg=COL["panel"],
                 fg=COL["ink"], font=self.f_title).pack(side="left", pady=17)
        tk.Label(header, text="ローカルLLM ・ 社外送信なし", bg=COL["panel"],
                 fg=COL["ink_soft"], font=self.f_small).pack(side="left", padx=10, pady=22)

        bar = tk.Frame(self.root, bg=COL["bg"])
        bar.pack(fill="x", padx=16, pady=(14, 6))
        self._btn(bar, "\U0001F4C1  フォルダを選択", self.choose_folder).pack(side="left")
        self._btn(bar, "\u27F3  読み込み", self.load_folder, primary=True).pack(side="left", padx=(8, 0))
        self.folder_var = tk.StringVar(value="フォルダ未選択")
        tk.Label(bar, textvariable=self.folder_var, bg=COL["bg"], fg=COL["ink_soft"],
                 font=self.f_small, anchor="w").pack(side="left", padx=12, fill="x", expand=True)
        tk.Label(bar, text="モデル", bg=COL["bg"], fg=COL["ink_soft"], font=self.f_small).pack(side="left")
        self.model_var = tk.StringVar(value=self.model_name)
        tk.Entry(bar, textvariable=self.model_var, width=18, font=self.f_small,
                 relief="flat", bg=COL["card"], fg=COL["ink"], highlightthickness=1,
                 highlightbackground=COL["line"], highlightcolor=COL["accent"]
                 ).pack(side="left", padx=(4, 0), ipady=4)

        chat_wrap = tk.Frame(self.root, bg=COL["bg"])
        chat_wrap.pack(fill="both", expand=True, padx=16, pady=8)
        self.chat = tk.Text(chat_wrap, wrap="word", font=self.f_body,
                            bg=COL["panel"], fg=COL["ink"], relief="flat",
                            padx=18, pady=16, spacing1=3, spacing3=8,
                            highlightthickness=1, highlightbackground=COL["line"],
                            insertbackground=COL["ink"], state="disabled")
        sb = tk.Scrollbar(chat_wrap, command=self.chat.yview, relief="flat",
                          troughcolor=COL["bg"], bg=COL["line"], width=12)
        self.chat.config(yscrollcommand=sb.set)
        sb.pack(side="right", fill="y")
        self.chat.pack(side="left", fill="both", expand=True)
        self.chat.tag_config("user_name", foreground=COL["accent_dk"], font=self.f_bold, spacing1=12)
        self.chat.tag_config("user_msg", foreground=COL["ink"], lmargin1=12, lmargin2=12)
        self.chat.tag_config("bot_name", foreground=COL["accent"], font=self.f_bold, spacing1=12)
        self.chat.tag_config("bot_msg", foreground=COL["ink"], lmargin1=12, lmargin2=12)
        self.chat.tag_config("sys", foreground=COL["ink_soft"], font=self.f_small, lmargin1=4, lmargin2=4)
        self.chat.tag_config("src", foreground=COL["green"], font=self.f_small, lmargin1=12, lmargin2=12, spacing1=4)

        input_wrap = tk.Frame(self.root, bg=COL["bg"])
        input_wrap.pack(fill="x", padx=16, pady=(4, 8))
        ebox = tk.Frame(input_wrap, bg=COL["card"], highlightthickness=1,
                        highlightbackground=COL["line"], highlightcolor=COL["accent"])
        ebox.pack(side="left", fill="x", expand=True, ipady=2)
        self.entry = tk.Entry(ebox, font=self.f_body, relief="flat", bg=COL["card"],
                              fg=COL["ink"], insertbackground=COL["accent"])
        self.entry.pack(fill="x", expand=True, padx=12, pady=8)
        self.entry.bind("<Return>", lambda e: self.on_send())

        # 送信ボタン(生成中は隠す)と停止ボタン(生成中だけ表示)を同じ場所に重ねる
        self.send_btn = self._btn(input_wrap, "送信  \u27A4", self.on_send, primary=True)
        self.send_btn.pack(side="left", padx=(8, 0))
        self.stop_btn = tk.Button(input_wrap, text="\u25A0 停止", command=self.on_stop,
                                  bg=COL["stop"], fg="#FFFFFF",
                                  activebackground=COL["stop_dk"], activeforeground="#FFFFFF",
                                  font=self.f_bold, relief="flat", bd=0,
                                  padx=16, pady=8, cursor="hand2")
        # stop_btn は最初は表示しない(on_sendで表示、完了/停止で隠す)

        self.status = tk.StringVar(value="\u2460 フォルダを選択 \u2192 \u2461 読み込み \u2192 \u2462 質問を入力")
        sbar = tk.Frame(self.root, bg=COL["panel"])
        sbar.pack(fill="x", side="bottom")
        tk.Frame(sbar, bg=COL["line"], height=1).pack(fill="x")
        tk.Label(sbar, textvariable=self.status, bg=COL["panel"], fg=COL["ink_soft"],
                 font=self.f_small, anchor="w", padx=16, pady=6).pack(fill="x")

        self._append("sys", "使い方:  \u2460「フォルダを選択」 \u2192 \u2461「読み込み」 \u2192 \u2462質問を入力して送信\n"
                            "(回答中は「停止」ボタンで中断できます。ログはターミナルに表示)\n")

    def _append(self, tag, text):
        self.chat.config(state="normal")
        self.chat.insert("end", text, tag)
        self.chat.see("end")
        self.chat.config(state="disabled")

    # 送信ボタン⇄停止ボタンの切り替え
    def _show_stop(self):
        self.send_btn.pack_forget()
        self.stop_btn.pack(side="left", padx=(8, 0))

    def _show_send(self):
        self.stop_btn.pack_forget()
        self.send_btn.pack(side="left", padx=(8, 0))
        self.send_btn.config(state="normal")

    def _poll_queue(self):
        try:
            while True:
                tag, text = self.msg_queue.get_nowait()
                if tag == "status":
                    self.status.set(text)
                elif tag == "show_send":
                    self._show_send()
                elif tag == "show_stop":
                    self._show_stop()
                elif tag == "error_box":
                    messagebox.showerror("エラー", text)
                else:
                    self._append(tag, text)
        except queue.Empty:
            pass
        self.root.after(80, self._poll_queue)

    def q(self, tag, text):
        self.msg_queue.put((tag, text))

    def choose_folder(self):
        path = filedialog.askdirectory(title="参照する文書フォルダを選択")
        if path:
            self.folder = Path(path)
            self.folder_var.set(str(self.folder))
            self.status.set("「読み込み」を押してインデックスを作成してください")
            log(f"フォルダ選択: {self.folder}")

    def load_folder(self):
        if not self.folder:
            self.status.set("先にフォルダを選択してください")
            log("フォルダ未選択のまま読み込みが押されました", "WARN")
            return
        if self.generating:
            self.status.set("回答生成中です。停止してから読み込んでください")
            return
        threading.Thread(target=self._load_worker, daemon=True).start()

    def _load_worker(self):
        try:
            log("=== 読み込み開始 ===")
            self.q("status", "ファイルを走査中...")
            files = [p for p in self.folder.rglob("*") if p.suffix.lower() in LOADERS]
            log(f"対象ファイル数: {len(files)}")
            if not files:
                self.q("sys", "[!] 対応ファイル(PDF/Word/Excel)が見つかりません\n")
                self.q("status", "対応ファイルなし")
                return

            raw = []
            for p in files:
                try:
                    loaded = LOADERS[p.suffix.lower()](p)
                    raw.extend(loaded)
                    log(f"読込OK: {p.name} ({len(loaded)}ブロック)")
                    self.q("sys", f"  読込: {p.name} ({len(loaded)}ブロック)\n")
                except Exception:
                    log_exc(f"ファイル読込 {p.name}")
                    self.q("sys", f"  [skip] {p.name} (ターミナルにエラー出力)\n")

            if not raw:
                self.q("sys", "[!] テキストを抽出できませんでした(スキャンPDFの可能性)\n")
                self.q("status", "抽出失敗")
                log("テキスト抽出ゼロ。スキャンPDFの可能性。", "WARN")
                return

            log("テキスト分割中...")
            self.q("status", "テキストを分割中...")
            from langchain_text_splitters import RecursiveCharacterTextSplitter
            from langchain_core.documents import Document as LCDocument
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=800, chunk_overlap=100,
                separators=["\n\n", "\n", "。", "、", " ", ""])
            lc_docs = []
            for d in raw:
                for chunk in splitter.split_text(d["text"]):
                    lc_docs.append(LCDocument(page_content=chunk,
                                              metadata={"source": d["source"], "loc": d["loc"]}))
            log(f"チャンク数: {len(lc_docs)}")

            log(f"埋め込みモデル準備中: {self.embed_model} (初回はDLに時間)")
            self.q("status", "埋め込みモデルを準備中(初回はDL)...")
            try:
                from langchain_huggingface import HuggingFaceEmbeddings
            except ImportError:
                log("langchain_huggingface未導入。community版にフォールバック", "WARN")
                from langchain_community.embeddings import HuggingFaceEmbeddings
            embeddings = HuggingFaceEmbeddings(
                model_name=self.embed_model,
                encode_kwargs={"normalize_embeddings": True})
            log("埋め込みモデル準備OK")

            log("ベクトル化中...")
            self.q("status", f"{len(lc_docs)}チャンクを埋め込み中...")
            from langchain_community.vectorstores import Chroma
            vectordb = Chroma.from_documents(documents=lc_docs, embedding=embeddings)
            self.retriever = vectordb.as_retriever(search_kwargs={"k": 4})
            log("ベクトルDB構築OK")

            self.model_name = self.model_var.get().strip() or "qwen3-32b:latest"
            log(f"LLM接続: {self.model_name}")
            from langchain_ollama import OllamaLLM
            self.llm = OllamaLLM(model=self.model_name, temperature=0.1)

            log("=== 読み込み完了 ===")
            self.q("sys", f"[準備完了] {len(files)}ファイル / {len(lc_docs)}チャンク。質問できます。\n")
            self.q("status", f"準備完了 ・ モデル: {self.model_name}")
        except Exception as e:
            log_exc("読み込み処理")
            self.q("sys", f"[エラー] {e}\n(詳細はターミナルを確認)\n")
            self.q("error_box", f"読み込み中にエラー:\n{e}\n\n詳細はターミナルを確認してください。")
            self.q("status", "読み込み失敗")

    def on_send(self):
        question = self.entry.get().strip()
        if not question:
            return
        if not self.retriever or not self.llm:
            self.status.set("先にフォルダを読み込んでください")
            log("未読み込みのまま送信が押されました", "WARN")
            return
        if self.generating:
            return
        self.entry.delete(0, "end")
        self._append("user_name", "\nあなた\n")
        self._append("user_msg", f"{question}\n")
        self.generating = True
        self.stop_flag.clear()
        self.q("show_stop", "")
        log(f"質問: {question}")
        threading.Thread(target=self._answer_worker, args=(question,), daemon=True).start()

    def on_stop(self):
        """停止ボタン: 生成中の出力を中断する。"""
        if self.generating:
            self.stop_flag.set()
            self.status.set("停止しています...")
            log("停止要求を受け付けました", "WARN")

    def _answer_worker(self, question):
        try:
            self.q("status", "関連箇所を検索中...")
            retrieved = self.retriever.invoke(question)
            log(f"検索ヒット: {len(retrieved)}件")
            blocks = []
            for i, doc in enumerate(retrieved, 1):
                src = doc.metadata.get("source", "不明")
                loc = doc.metadata.get("loc", "")
                blocks.append(f"[資料{i}] (出典: {src} {loc})\n{doc.page_content}")
            context = "\n\n".join(blocks)
            prompt = ANSWER_TEMPLATE.format(system=SYSTEM_PROMPT, context=context, question=question)

            self.q("status", "回答を生成中...(停止ボタンで中断可)")
            self.q("bot_name", "\nアシスタント\n")
            log("LLM応答生成中...")

            stopped = False
            for chunk in self.llm.stream(prompt):
                if self.stop_flag.is_set():
                    stopped = True
                    self.q("bot_msg", "\n\n[停止しました]")
                    log("ユーザー操作により生成を停止")
                    break
                self.q("bot_msg", chunk)

            if not stopped:
                used = sorted({f"{d.metadata.get('source','?')} {d.metadata.get('loc','')}".strip()
                               for d in retrieved})
                self.q("src", "\n\U0001F4CE 参照したファイル: " + " / ".join(used) + "\n")
                log("回答完了")
            self.q("status", f"準備完了 ・ モデル: {self.model_name}")
        except Exception as e:
            log_exc("回答生成")
            self.q("sys", f"\n[エラー] {e}\n(詳細はターミナルを確認)\n")
            self.q("error_box", f"回答中にエラー:\n{e}\n\n詳細はターミナルを確認してください。")
            self.q("status", "回答失敗")
        finally:
            self.generating = False
            self.stop_flag.clear()
            self.q("show_send", "")


def main():
    log("=== rag_chat 起動 ===")
    log(f"Python: {sys.version.split()[0]}")
    try:
        root = tk.Tk()
        RagApp(root)
        root.mainloop()
        log("=== 正常終了 ===")
    except Exception:
        log_exc("アプリ起動")
        input("\nエラーが発生しました。Enterキーで閉じます...")


if __name__ == "__main__":
    main()
